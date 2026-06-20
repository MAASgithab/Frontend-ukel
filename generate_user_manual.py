from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============ STYLE CONFIG ============
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============ HELPER FUNCTIONS ============
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x06, 0x65, 0xFE)
    return h

def add_paragraph_bold_first(text_bold, text_normal=""):
    p = doc.add_paragraph()
    run_bold = p.add_run(text_bold)
    run_bold.bold = True
    if text_normal:
        p.add_run(text_normal)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def add_step(number, text):
    p = doc.add_paragraph()
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    p.add_run(text)
    return p

# ============ COVER PAGE ============
for _ in range(6):
    doc.add_paragraph()

p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_cover.add_run("USER MANUAL WEB")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x06, 0x65, 0xFE)

p_cover2 = doc.add_paragraph()
p_cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_cover2.add_run("BATIKNESIA")
run2.bold = True
run2.font.size = Pt(36)
run2.font.color.rgb = RGBColor(0x06, 0xFE, 0x9F)

doc.add_paragraph()

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Panduan Menggunakan Website Batiknesia\nuntuk Belanja Kain Tradisional Indonesia")
run_sub.font.size = Pt(14)
run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(4):
    doc.add_paragraph()

p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_name = p_name.add_run("Muhammad Abdurrahim As-Silmi\n12410486\nTajur 1\nTugas Proyek React Mandiri")
run_name.font.size = Pt(12)
run_name.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ============ DAFTAR ISI ============
add_heading_styled("Daftar Isi", level=1)

toc_items = [
    ("Kata Pengantar", ""),
    ("Pengenalan", ""),
    ("User Aplikasi", ""),
    ("  1. Peran Akun User", ""),
    ("  2. Peran Akun Admin", ""),
    ("Alur Penggunaan Website Batiknesia", ""),
    ("Interface Aplikasi", ""),
    ("  1. Halaman Awal (Home)", ""),
    ("  2. Halaman Register dan Login", ""),
    ("    2.1. Register", ""),
    ("    2.2. Login", ""),
    ("    2.3. Logout", ""),
    ("  3. Halaman User", ""),
    ("    3.1. Melihat Daftar Produk", ""),
    ("    3.2. Mencari Produk", ""),
    ("    3.3. Menambahkan Produk ke Keranjang", ""),
    ("    3.4. Melihat Keranjang Belanja", ""),
    ("    3.5. Membuat Pesanan (Order)", ""),
    ("  4. Halaman Admin", ""),
    ("    4.1. Kelola Produk (Kain)", ""),
    ("    4.2. Menambah Produk Baru", ""),
    ("    4.3. Mengedit Produk", ""),
    ("    4.4. Menghapus Produk", ""),
    ("    4.5. Kelola Order", ""),
    ("    4.6. Mengubah Status Order", ""),
    ("    4.7. Menghapus Order", ""),
    ("    4.8. Kelola User", ""),
    ("    4.9. Mengedit Data User", ""),
    ("    4.10. Menghapus User", ""),
]

for item, _ in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ============ KATA PENGANTAR ============
add_heading_styled("Kata Pengantar", level=1)

kata_pengantar = (
    "Puji dan syukur kami panjatkan kepada Tuhan Yang Maha Esa atas segala rahmat "
    "dan karunia-Nya sehingga website e-commerce kain tradisional Batiknesia ini "
    "dapat dikembangkan dengan baik.\n\n"
    "Website Batiknesia ini dibuat untuk membantu seluruh pengguna dalam membeli "
    "dan menjelajahi berbagai kain tradisional asli Indonesia secara mudah dan "
    "terstruktur. Batiknesia hadir sebagai platform digital yang menghubungkan "
    "pengrajin kain tradisional langsung kepada pelanggan, mendukung pelestarian "
    "budaya dan UMKM Nusantara.\n\n"
    "Website Batiknesia digunakan untuk menyajikan berbagai produk kain tradisional "
    "seperti Batik, Tenun, Songket, Ulos, Lurik, dan Kain Nusantara lainnya. "
    "Pengguna dapat melihat produk, menambahkan ke keranjang, dan melakukan "
    "pemesanan dengan mudah.\n\n"
    "Diharapkan panduan ini dapat menjadi acuan praktis dalam menggunakan setiap "
    "fitur yang tersedia di website Batiknesia.\n\n"
    "Semoga website Batiknesia ini dapat digunakan dengan sebaik-baiknya dan "
    "memberikan manfaat bagi seluruh pengguna dalam melestarikan budaya Indonesia."
)
doc.add_paragraph(kata_pengantar)

doc.add_page_break()

# ============ PENGENALAN ============
add_heading_styled("Pengenalan", level=1)

pengenalan = (
    "Selamat datang di buku panduan untuk Batiknesia! Batiknesia adalah website "
    "e-commerce sederhana yang dibuat khusus untuk membantu kalian membeli dan "
    "menjelajahi kain tradisional Indonesia dengan mudah.\n\n"
    "Dengan Batiknesia, kalian bisa:\n"
    "- Melihat berbagai produk kain tradisional asli dari pengrajin Nusantara\n"
    "- Mencari produk berdasarkan nama kain\n"
    "- Menambahkan produk ke keranjang belanja\n"
    "- Melakukan pemesanan dengan mudah\n"
    "- Admin dapat mengelola produk, order, dan user\n\n"
    "Buku panduan ini akan menjelaskan langkah demi langkah cara menggunakan semua "
    "fitur Batiknesia. Yuk, mulai belajar dan lestarikan budaya Indonesia!"
)
doc.add_paragraph(pengenalan)

doc.add_page_break()

# ============ USER APLIKASI ============
add_heading_styled("User Aplikasi", level=1)

doc.add_paragraph(
    "Pada web Batiknesia ada dua jenis peran yaitu User dan Admin."
)

add_paragraph_bold_first("1. Peran Akun User\n",
    "Akun user adalah akun untuk para pengguna biasa yang ingin membeli "
    "kain tradisional. Disini mereka bisa melihat produk, mencari produk, "
    "menambahkan ke keranjang, dan melakukan pemesanan."
)

add_paragraph_bold_first("2. Peran Akun Admin\n",
    "Akun Admin digunakan oleh operator untuk mengelola website. Admin dapat "
    "mengelola produk (tambah, edit, hapus kain), mengelola order (mengubah "
    "status, menghapus order), dan mengelola user (edit, hapus user)."
)

doc.add_page_break()

# ============ ALUR PENGGUNAAN ============
add_heading_styled("Alur Penggunaan Website Batiknesia", level=1)

alur_items = [
    "1. Pengguna membuka halaman utama website Batiknesia. Terdapat menu Beranda, Produk, Tentang Kami, Login, dan Daftar. Halaman ini menjadi pintu masuk pertama bagi semua pengunjung.",
    "2. Pengunjung yang sudah punya akun bisa login ke halaman utama. Di sini mereka bisa melihat produk, mencari produk, dan menambahkan ke keranjang.",
    "3. Pengunjung yang belum punya akun bisa mendaftar melalui halaman Register.",
    "4. Setelah login, user bisa menambahkan produk ke keranjang dan melanjutkan ke halaman order untuk membuat pesanan.",
    "5. Admin login ke dashboard khusus dengan hak akses lebih tinggi. Admin bisa mengelola produk, order, dan data semua user yang terdaftar.",
]

for item in alur_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ============ INTERFACE APLIKASI ============
add_heading_styled("Interface Aplikasi", level=1)
doc.add_paragraph("Ini bagian memperjelas tampilan halaman web Batiknesia.")

# --- 1. Halaman Awal ---
add_heading_styled("1. Halaman Awal (Home)", level=2)
doc.add_paragraph(
    "Halaman Awal atau Home digunakan untuk menyambut pengguna. Di halaman ini "
    "terdapat:\n"
    "- Navbar dengan menu Beranda, Produk, Tentang Kami, serta tombol Login dan Daftar\n"
    "- Hero Section dengan gambar kain tradisional dan ajakan belanja\n"
    "- Banner Promo diskon hingga 20%\n"
    "- Keunggulan Batiknesia (produk asli, pengiriman nasional, pembayaran aman, dukung UMKM)\n"
    "- Tentang Batiknesia dengan informasi visi dan misi\n"
    "- Footer dengan kontak dan media sosial"
)

doc.add_page_break()

# --- 2. Register dan Login ---
add_heading_styled("2. Halaman Register dan Login", level=2)
doc.add_paragraph("Halaman ini untuk pengguna agar bisa masuk ke website.")

add_heading_styled("2.1. Register", level=3)
register_steps = [
    "Masuk terlebih dahulu ke website Batiknesia",
    "Tekan tombol \"Daftar\" di pojok kanan navbar",
    "Masukan Nama Lengkap, Email, Password (minimal 6 karakter)",
    "Isi No. Telepon dan Alamat (opsional)",
    "Tekan tombol \"Daftar\"",
    "Setelah berhasil, akan diarahkan ke halaman Login",
]
for i, step in enumerate(register_steps, 1):
    add_step(i, step)

add_heading_styled("2.2. Login", level=3)
login_steps = [
    "Masuk ke halaman Login dengan menekan tombol \"Login\" di navbar",
    "Masukan email dan password yang sudah didaftarkan",
    "Tekan tombol \"Masuk\"",
    "Jika berhasil, akan masuk ke halaman utama",
    "Jika role admin, akan diarahkan ke halaman Admin",
]
for i, step in enumerate(login_steps, 1):
    add_step(i, step)

add_heading_styled("2.3. Logout", level=3)
logout_steps = [
    "Setelah login, nama user akan muncul di navbar",
    "Klik tombol \"Logout\" di sebelah nama user",
    "User akan keluar dari akun dan kembali ke halaman utama",
]
for i, step in enumerate(logout_steps, 1):
    add_step(i, step)

doc.add_page_break()

# --- 3. Halaman User ---
add_heading_styled("3. Halaman User", level=2)
doc.add_paragraph("Disini, User bisa mengakses fitur-fitur website.")

add_heading_styled("3.1. Melihat Daftar Produk", level=3)
produk_steps = [
    "Klik menu \"Produk\" di navbar",
    "Halaman akan menampilkan daftar produk kain tradisional",
    "Setiap produk menampilkan gambar, kategori, nama, lebar kain, dan harga",
    "Tersedia tombol \"Detail\" dan \"Keranjang\" untuk setiap produk",
    "Produk ditampilkan dengan sistem paginasi (8 produk per halaman)",
]
for i, step in enumerate(produk_steps, 1):
    add_step(i, step)

add_heading_styled("3.2. Mencari Produk", level=3)
search_steps = [
    "Di halaman Produk, terdapat kolom pencarian di bagian atas",
    "Ketikkan nama produk yang ingin dicari",
    "Produk akan otomatis difilter sesuai kata kunci yang diketikkan",
]
for i, step in enumerate(search_steps, 1):
    add_step(i, step)

add_heading_styled("3.3. Menambahkan Produk ke Keranjang", level=3)
cart_steps = [
    "Cari produk yang diinginkan",
    "Klik tombol \"Keranjang\" pada produk tersebut",
    "Akan muncul notifikasi bahwa produk berhasil ditambahkan",
    "Jumlah di ikon keranjang (navbar) akan bertambah",
]
for i, step in enumerate(cart_steps, 1):
    add_step(i, step)

add_heading_styled("3.4. Melihat Keranjang Belanja", level=3)
view_cart_steps = [
    "Klik ikon keranjang di pojok kanan navbar",
    "Halaman keranjang menampilkan semua produk yang ditambahkan",
    "User bisa menambah atau mengurangi jumlah produk dengan tombol + dan -",
    "User bisa menghapus produk dari keranjang dengan tombol hapus",
    "Subtotal setiap produk dan total keseluruhan ditampilkan",
    "Klik tombol \"Beli\" untuk melanjutkan ke halaman order",
]
for i, step in enumerate(view_cart_steps, 1):
    add_step(i, step)

add_heading_styled("3.5. Membuat Pesanan (Order)", level=3)
order_steps = [
    "Di halaman Detail Pesanan, periksa kembali daftar produk yang akan dibeli",
    "Struk pesanan menampilkan nama toko (BATIKNESIA), tanggal, daftar item, dan total harga",
    "Klik tombol \"Konfirmasi Pesanan\" untuk membuat pesanan",
    "Pesanan akan dikirim ke sistem dan keranjang akan dikosongkan",
    "Halaman sukses akan muncul dengan ID Pesanan",
    "Klik \"Kembali ke Beranda\" untuk kembali ke halaman utama",
]
for i, step in enumerate(order_steps, 1):
    add_step(i, step)

doc.add_page_break()

# --- 4. Halaman Admin ---
add_heading_styled("4. Halaman Admin", level=2)
doc.add_paragraph(
    "Disini halaman hanya bisa diakses oleh Admin saja. Admin harus login "
    "dengan akun yang memiliki role admin. Setelah login, admin akan otomatis "
    "diarahkan ke halaman Admin."
)

add_heading_styled("4.1. Kelola Produk (Kain)", level=3)
doc.add_paragraph(
    "Di halaman Admin, tab pertama adalah \"Kelola Produk\". Disini admin "
    "dapat melihat daftar semua produk kain yang tersedia."
)

add_heading_styled("4.2. Menambah Produk Baru", level=3)
tambah_produk_steps = [
    "Klik tombol \"Tambah Produk\" di bagian atas tabel",
    "Akan muncul modal form tambah produk",
    "Isi Nama Produk, pilih Kategori (Batik, Tenun, Songket, Ulos, Lurik, Kain Nusantara)",
    "Masukkan Harga (Rp), Lebar (cm), dan Stok",
    "Isi Deskripsi produk",
    "Upload Gambar produk (format gambar)",
    "Klik tombol \"Tambah\" untuk menyimpan produk",
    "Produk baru akan muncul di tabel dan halaman produk user",
]
for i, step in enumerate(tambah_produk_steps, 1):
    add_step(i, step)

add_heading_styled("4.3. Mengedit Produk", level=3)
edit_produk_steps = [
    "Di tabel produk, klik ikon pensil (edit) pada produk yang ingin diubah",
    "Akan muncul modal form dengan data produk yang sudah terisi",
    "Ubah data yang diperlukan (nama, kategori, harga, lebar, stok, deskripsi)",
    "Jika ingin mengganti gambar, upload gambar baru",
    "Klik tombol \"Simpan\" untuk menyimpan perubahan",
]
for i, step in enumerate(edit_produk_steps, 1):
    add_step(i, step)

add_heading_styled("4.4. Menghapus Produk", level=3)
hapus_produk_steps = [
    "Di tabel produk, klik ikon tempat sampah (delete) pada produk yang ingin dihapus",
    "Akan muncul konfirmasi \"Apakah Anda yakin ingin menghapus produk ini?\"",
    "Klik \"OK\" untuk menghapus, atau \"Batal\" untuk membatalkan",
    "Produk akan langsung terhapus dari tabel dan halaman produk user",
]
for i, step in enumerate(hapus_produk_steps, 1):
    add_step(i, step)

add_heading_styled("4.5. Kelola Order", level=3)
doc.add_paragraph(
    "Klik tab \"Kelola Order\" di sidebar admin untuk melihat daftar semua "
    "pesanan yang masuk dari user."
)

add_heading_styled("4.6. Mengubah Status Order", level=3)
status_order_steps = [
    "Di tabel order, lihat kolom Status untuk mengetahui status pesanan saat ini",
    "Status berurutan: pending → diproses → dikirim → selesai → dibatalkan",
    "Klik ikon pensil (edit) untuk mengubah status ke status berikutnya",
    "Status akan berubah sesuai urutan (misal: pending jadi diproses)",
    "Setiap status memiliki warna berbeda: hijau (selesai), merah (dibatalkan), biru (dikirim), kuning (diproses), abu-abu (pending)",
]
for i, step in enumerate(status_order_steps, 1):
    add_step(i, step)

add_heading_styled("4.7. Menghapus Order", level=3)
hapus_order_steps = [
    "Di tabel order, klik ikon tempat sampah (delete) pada order yang ingin dihapus",
    "Akan muncul konfirmasi \"Apakah Anda yakin ingin menghapus order ini?\"",
    "Klik \"OK\" untuk menghapus",
    "Order akan langsung terhapus dari daftar",
]
for i, step in enumerate(hapus_order_steps, 1):
    add_step(i, step)

add_heading_styled("4.8. Kelola User", level=3)
doc.add_paragraph(
    "Klik tab \"Kelola User\" di sidebar admin untuk melihat daftar semua "
    "user yang terdaftar di website."
)

add_heading_styled("4.9. Mengedit Data User", level=3)
edit_user_steps = [
    "Di tabel user, klik ikon pensil (edit) pada user yang ingin diubah",
    "Akan muncul modal form edit user",
    "Ubah data yang diperlukan: Nama Lengkap, Email, No. Telepon, Alamat",
    "Ubah Role jika perlu (Admin atau User)",
    "Klik tombol \"Simpan\" untuk menyimpan perubahan",
]
for i, step in enumerate(edit_user_steps, 1):
    add_step(i, step)

add_heading_styled("4.10. Menghapus User", level=3)
hapus_user_steps = [
    "Di tabel user, klik ikon tempat sampah (delete) pada user yang ingin dihapus",
    "Akan muncul konfirmasi \"Apakah Anda yakin ingin menghapus user ini?\"",
    "Klik \"OK\" untuk menghapus",
    "User akan langsung terhapus dari daftar",
]
for i, step in enumerate(hapus_user_steps, 1):
    add_step(i, step)

# ============ SAVE ============
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "User Manual Web Batiknesia.docx")
doc.save(output_path)
print(f"File berhasil disimpan di: {output_path}")